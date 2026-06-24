"""
Document parsing service for extracting text from various file formats.
"""

import io
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional
import time

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.ai.models import DocumentContent


class DocumentParser:
    """Service for parsing documents and extracting text content."""
    
    def __init__(self):
        self.supported_types = []
        if PDF_AVAILABLE:
            self.supported_types.extend(['.pdf'])
        if DOCX_AVAILABLE:
            self.supported_types.extend(['.docx', '.doc'])
        
        # Always support plain text
        self.supported_types.extend(['.txt', '.md'])
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file type is supported for parsing."""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_types
    
    def parse_document(self, file_path: str, mime_type: Optional[str] = None) -> DocumentContent:
        """Parse document and extract text content."""
        start_time = time.time()
        
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.pdf':
                return self._parse_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            # Return error document content
            return DocumentContent(
                text=f"Error parsing document: {str(e)}",
                metadata={
                    "error": str(e),
                    "file_path": file_path,
                    "processing_time": time.time() - start_time
                },
                file_type=Path(file_path).suffix.lower(),
                extraction_method="error"
            )
    
    def _parse_pdf(self, file_path: str) -> DocumentContent:
        """Parse PDF document."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        text_content = []
        page_count = 0
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        text_content.append(f"--- Page {page_num + 1} (Error) ---\nError extracting text: {str(e)}")
        
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
        
        return DocumentContent(
            text="\n\n".join(text_content),
            metadata={
                "page_count": page_count,
                "file_path": file_path,
                "total_pages": page_count
            },
            page_count=page_count,
            file_type=".pdf",
            extraction_method="PyPDF2"
        )
    
    def _parse_docx(self, file_path: str) -> DocumentContent:
        """Parse DOCX document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine content
            content_parts = []
            if paragraphs:
                content_parts.append("--- Document Text ---\n" + "\n\n".join(paragraphs))
            if tables_text:
                content_parts.append("--- Tables ---\n" + "\n\n".join(tables_text))
            
            return DocumentContent(
                text="\n\n".join(content_parts),
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    "file_path": file_path
                },
                file_type=".docx",
                extraction_method="python-docx"
            )
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_text(self, file_path: str) -> DocumentContent:
        """Parse plain text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            line_count = len(content.split('\n'))
            
            return DocumentContent(
                text=content,
                metadata={
                    "line_count": line_count,
                    "character_count": len(content),
                    "file_path": file_path
                },
                file_type=Path(file_path).suffix.lower(),
                extraction_method="text_reader"
            )
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    return DocumentContent(
                        text=content,
                        metadata={
                            "encoding_used": encoding,
                            "file_path": file_path
                        },
                        file_type=Path(file_path).suffix.lower(),
                        extraction_method=f"text_reader_{encoding}"
                    )
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any supported encoding")
        
        except Exception as e:
            raise Exception(f"Failed to parse text file: {str(e)}")


# Global parser instance
document_parser = DocumentParser()